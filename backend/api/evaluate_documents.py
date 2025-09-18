from fastapi import Request, UploadFile, File, WebSocket, WebSocketDisconnect, APIRouter, BackgroundTasks, Form
from fastapi.middleware.cors import CORSMiddleware
import json
from openai import OpenAI
from docx import Document
from openpyxl import Workbook
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import os
import tempfile
from gtts import gTTS
import base64 
from pptx import Presentation
import fitz  # PyMuPDF
import os
import shutil
from typing import List, Optional
from fastapi.responses import JSONResponse
from fastapi_mail import FastMail, MessageSchema, ConnectionConfig
# import whisper
import base64
from utils.doc_utils import extract_text_pdf, extract_text_ppt
from utils.audio_utils import get_audio_base64
from utils.session_utils import get_session, save_session, session_file
from models.gemini_client import GeminiClient

# whisper_model = whisper.load_model("base")

USE_GPT_STT = False

language = 'en'
geminiClient = GeminiClient()  # Initialize GeminiClient when needed

# -----------------------------
# Configure FastAPI-Mail
# -----------------------------
# conf = ConnectionConfig(
#     MAIL_USERNAME="Tharun Ganapathi",
#     MAIL_PASSWORD="Gana@1869",
#     MAIL_FROM="tharunganapathi19@gmail.com",
#     MAIL_PORT=587,
#     MAIL_SERVER="smtp.gmail.com",
#     MAIL_TLS=True,
#     MAIL_SSL=False,
#     USE_CREDENTIALS=True,
#     VALIDATE_CERTS=True
# )

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

SESSION_DIR = "sessions"
os.makedirs(SESSION_DIR, exist_ok=True)

router = APIRouter()

# Memory of collected info per session
conversations = {}

# ---------- WebSocket voice streaming ----------
@router.post("/chat/{session_id}")
async def chat(session_id: str, request: Request):
    body = await request.json()
    user_msg = body.get("message")
    custom_prompt = body.get("prompt", '''
# 🎤 Startup Founder Voice Intake Assistant Prompt

You are a **friendly, professional voice intake assistant**. Your goal is to guide a startup founder through a conversation, collect structured details, and make them feel comfortable sharing their story.  

---

## Conversation Rules
1. Speak **conversationally** — no forms or surveys. Use natural, flowing language.  
2. Ask **one question at a time**. Wait for the founder’s spoken response before continuing.  
3. If they answer partially, ask a **follow-up or clarifying question** — don’t repeat the full original question.  
4. Encourage **storytelling**. If answers are short, politely nudge for more detail.  
5. Stay **human-like**: vary phrasing, use small acknowledgments (“Got it”, “That makes sense”, “Thanks for sharing”).  
6. After completing each section, **briefly summarize** what you captured and confirm with them.  
7. Keep responses **concise, warm, and spoken-friendly**.  

---

## Goal
Capture these structured fields:  
- **Founder & Idea Details**  
- **Market & Problem Exploration**  
- **Product & Solution Details**  
- **Operations & Team**  
- **Financials**  
- **Go-to-Market & Plan**  
- **Additional Info / Notes**  

---

## Section Prompts

### Opening / Introduction
- “Hi, thanks for taking the time to chat today. I’d love to learn more about you and your startup idea. Let’s start simple — could you tell me your name and how I can best reach you?”  
- “Great, thanks [Name]. And what’s the name of your business or startup?”  

### Founder & Idea Details
- “How would you describe your business idea in your own words?”  
- “What makes your idea stand out from what’s already out there?”  

### Market & Problem Exploration
- “Who do you imagine will benefit most from your product or service?”  
- “What’s the key problem or need you’re solving for them?”  
- “Do you have a sense of how big the market might be, or who your main competitors are?”  

### Product & Solution Details
- “Can you walk me through what your product or service actually does?”  
- “If you had to launch a simple first version, what features would it absolutely need?”  
- “What technologies or platforms do you think you’ll need to build it?”  

### Operations & Team
- “Tell me a little about your background — what led you to start this?”  
- “Do you have a team in mind already, or are you planning to bring in partners or advisors?”  

### Financials
- “What’s your rough budget estimate for getting this off the ground?”  
- “How much are you personally investing, and how much do you think you’ll need to raise?”  
- “What’s your plan for generating revenue? Subscriptions, one-time sales, something else?”  

### Go-to-Market & Plan
- “How are you thinking about launching? Soft launch, pilot, or straight to market?”  
- “What channels do you think will help you get your first customers?”  
- “Do you have milestones you want to hit in the first 6 to 12 months?”  

### Closing / Additional Notes
- “That was really helpful, thank you. Before we wrap up, is there anything else — like legal considerations, compliance issues, or special notes — that you’d like me to capture?”  
- “Perfect, thanks so much. I’ll compile everything we discussed and share it with you in a clear summary.”  

---

## Extra Realism Tips
- Vary acknowledgment phrases: “Got it”, “I see”, “That’s helpful”, “Interesting”, “Makes sense”.  
- Add **short pauses** in speech to sound natural.  
- End each section with a **mini-summary**: “So far, I have that your idea is X, targeting Y, with competitors Z. Does that sound right?”  
- If a question is skipped, move on gracefully: “No worries, we can come back to that later.”  

''')

    
    conv = get_session(session_id)

    # Store user message
    conv["messages"].append({"role": "user", "content": user_msg})

    # Call LLM (OpenAI example – swap model if using Mistral/LLaMA)
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": custom_prompt},
            *conv["messages"]
        ]
    )

    reply = response.choices[0].message.content
    conv["messages"].append({"role": "assistant", "content": reply})

    save_session(session_id, conv)

    speech = gTTS(text=reply, lang=language, slow=False)
    speech.save("output.mp3")

    audio_b64 = get_audio_base64()

    return {"reply": reply, "data": conv["data"], "audio_base64": audio_b64}


@router.get("/export/{session_id}/{format}")
async def export_data(session_id: str, format: str):
    
    data = get_session(session_id)
    print(data)
    if not data:
        return {"error": "No data available to export."}
    if format == "json":
        filename = "startup_details.json"
        with open(filename, "w") as f:
            json.dump(data, f, indent=4)
        return {"file": filename}

    elif format == "excel":
        filename = "startup_details.xlsx"
        wb = Workbook()
        ws = wb.active
        for section, details in data.items():
            ws.append([section])
            for k, v in details.items():
                ws.append([k, v])
            ws.append([])
        wb.save(filename)
        return {"file": filename}

    elif format == "word":
        filename = "startup_details.docx"
        doc = Document()
        for section, details in data.items():
            doc.add_heading(section, level=1)
            for k, v in details.items():
                doc.add_paragraph(f"{k}: {v}")
        doc.save(filename)
        return {"file": filename}

    elif format == "pdf":
        filename = "startup_details.pdf"
        doc = SimpleDocTemplate(filename)
        styles = getSampleStyleSheet()
        flowables = []
        for section, details in data.items():
            flowables.append(Paragraph(f"<b>{section}</b>", styles["Heading1"]))
            for k, v in details.items():
                flowables.append(Paragraph(f"{k}: {v}", styles["Normal"]))
        doc.build(flowables)
        return {"file": filename}

    return {"error": "Invalid format"}


def evaluate_startup_documents(uploaded_files, founder_name, founder_email, startup_name, request_id):
    # texts = ["Not Available", "Not Available", "Not applicable", "Not applicable"]
    docs = dict()
    for topic, file_path in uploaded_files.items():
        if file_path.endswith(".pdf"):
            text = extract_text_pdf(file_path)
        elif file_path.endswith((".ppt", ".pptx")):
            text = extract_text_ppt(file_path)
        docs[topic] = text
    
    summary = geminiClient.analyze_documents(request_id=request_id,founder_name=founder_name,founder_email=founder_email,startup_name=startup_name,docs=docs)

    # combined_text = "\n\n".join(texts)

#     prompt = f"""
# You are analyzing startup founder documents (pitch deck & checklist). Provide the following template completely filled with necessary details grounded as per documents data:


# # COMPREHENSIVE INVESTOR-READY BUSINESS PLAN

# Request ID: {request_id}
# Founder: {founder_name} ({founder_email})
# Startup: {startup_name}

# ## EXECUTIVE SUMMARY
# *[This section should be written last but appears first - keep to 1-2 pages maximum]*

# **Company Overview:**
# - Company Name: [Collected Input]
# - Mission Statement: [Collected Input]
# - Vision Statement: [Collected Input]
# - One-Line Value Proposition: [Collected Input]

# **The Opportunity:**
# - Problem Statement: [Brief description of the customer problem/need]
# - Market Size: [Total Addressable Market (TAM), Serviceable Addressable Market (SAM), Serviceable Obtainable Market (SOM)]
# - Unique Solution: [How your solution uniquely addresses the problem]

# **Financial Highlights:**
# - Funding Amount Requested: [Collected Input]
# - Use of Funds: [Key categories where money will be invested]
# - Revenue Projections: [Year 1, Year 2, Year 3 projected revenues]
# - Expected ROI for Investors: [Projected return multiple and timeline]

# **Investment Highlights:**
# - Key traction metrics or early validation
# - Competitive advantages
# - Team credentials
# - Clear path to profitability

# ---

# ## 1. COMPANY OVERVIEW & CONCEPT

# ### 1.1 Business Concept
# **Business Name:** [Collected Input]
# **Legal Structure:** [Corporation/LLC/Partnership - to be determined]
# **Industry:** [Primary industry classification]
# **Business Model Type:** [B2B/B2C/B2B2C/Marketplace/SaaS/etc.]

# ### 1.2 Mission & Vision
# **Mission Statement:** [Collected Input - What the company does and why it exists]
# **Vision Statement:** [Collected Input - Where the company is heading long-term]
# **Core Values:** [3-5 fundamental principles that guide the business]

# ### 1.3 Business Idea Summary
# [Collected Input - Expanded with:]
# - What specific problem does this solve?
# - How is this solution different from existing alternatives?
# - Why is now the right time for this solution?
# - What makes this scalable and defensible?

# ### 1.4 Unique Value Proposition
# [Collected Input - Enhanced with:]
# - Clear benefit statement
# - Differentiation from competitors
# - Quantifiable value delivered to customers
# - Emotional and rational reasons customers will choose you

# ---

# ## 2. MARKET OPPORTUNITY & ANALYSIS

# ### 2.1 Problem Statement
# **Customer Problem/Need:** [Collected Input - Expanded with:]
# - How big is this problem? (quantify with data)
# - How are customers currently solving this problem?
# - What are the costs/inefficiencies of current solutions?
# - How urgent is solving this problem for customers?

# ### 2.2 Target Market Analysis
# **Primary Target Market:** [Collected Input - Enhanced with:]
# - Detailed customer demographics/firmographics
# - Customer personas (3-5 detailed profiles)
# - Geographic markets (primary and expansion)
# - Customer acquisition difficulty and cost

# **Market Size & Opportunity:**
# - Total Addressable Market (TAM): [Collected Input - Enhanced with source]
# - Serviceable Addressable Market (SAM): [Realistic portion you can capture]
# - Serviceable Obtainable Market (SOM): [What you can realistically achieve in 3-5 years]
# - Market growth rate and trends

# ### 2.3 Competitive Landscape
# **Main Competition:** [Collected Input - Expanded with:]

# **Direct Competitors:**
# - Company A: [Products, market share, strengths, weaknesses]
# - Company B: [Products, market share, strengths, weaknesses]
# - Company C: [Products, market share, strengths, weaknesses]

# **Indirect Competitors:**
# - Alternative solutions customers currently use
# - Substitute products or services

# **Competitive Advantages:**
# - Technology advantages
# - Cost advantages
# - Network effects
# - Brand/reputation
# - Regulatory barriers
# - Switching costs
# - Patents/IP protection

# **Competitive Matrix:**
# | Feature/Capability | Your Company | Competitor A | Competitor B | Competitor C |
# |-------------------|--------------|-------------|-------------|-------------|
# | [Key Feature 1] | [Rating] | [Rating] | [Rating] | [Rating] |
# | [Key Feature 2] | [Rating] | [Rating] | [Rating] | [Rating] |
# | Price Point | [Your Price] | [Price] | [Price] | [Price] |

# ---

# ## 3. PRODUCT/SERVICE SOLUTION

# ### 3.1 Product/Service Description
# [Collected Input - Enhanced with:]
# - Detailed feature breakdown
# - Technical specifications (if applicable)
# - Service delivery model
# - Quality standards and certifications
# - Intellectual property protection

# ### 3.2 Core Features (MVP)
# [Collected Input - Organized by priority:]

# **Phase 1 Features (MVP):**
# 1. [Essential feature 1] - [Customer benefit]
# 2. [Essential feature 2] - [Customer benefit]
# 3. [Essential feature 3] - [Customer benefit]

# **Phase 2 Features (Post-MVP):**
# 1. [Advanced feature 1] - [Customer benefit]
# 2. [Advanced feature 2] - [Customer benefit]

# **Future Roadmap:**
# - Year 1: [Key developments]
# - Year 2: [Key developments]
# - Year 3+: [Vision for product evolution]

# ### 3.3 Technology & Development
# **Technologies/Platforms Needed:** [Collected Input - Enhanced with:]
# - Core technology stack
# - Infrastructure requirements
# - Third-party integrations
# - Security and compliance requirements
# - Scalability considerations

# **Development Timeline:** [Collected Input - Enhanced with:]
# - MVP completion: [Date and milestones]
# - Beta testing: [Timeline and metrics]
# - Full product launch: [Date and success criteria]
# - Major version updates: [Timeline]

# ### 3.4 Product Validation
# **Customer Discovery:**
# - Number of customer interviews conducted: [Number]
# - Key insights from customer feedback
# - Product-market fit indicators
# - Early customer commitments or pre-orders

# **Testing & Prototyping:**
# - Prototype development status
# - Beta testing results and feedback
# - Key performance metrics from testing
# - Product iteration based on feedback

# ---

# ## 4. BUSINESS MODEL & REVENUE STRATEGY

# ### 4.1 Revenue Model
# [Collected Input - Enhanced with:]

# **Primary Revenue Streams:**
# 1. [Revenue Stream 1]: [Description, pricing, % of total revenue]
# 2. [Revenue Stream 2]: [Description, pricing, % of total revenue]
# 3. [Revenue Stream 3]: [Description, pricing, % of total revenue]

# **Revenue Model Type:**
# - [ ] Subscription/SaaS
# - [ ] Transaction-based
# - [ ] Product sales
# - [ ] Service fees
# - [ ] Advertising/Marketplace
# - [ ] Licensing
# - [ ] Freemium
# - [ ] Other: [Specify]

# ### 4.2 Pricing Strategy
# [Collected Input - Enhanced with:]
# - Pricing methodology (cost-plus, value-based, competitive)
# - Price points for different customer segments
# - Pricing elasticity analysis
# - Discount and promotional strategies
# - Price optimization plans

# **Pricing Tiers/Options:**
# | Tier | Target Customer | Price Point | Key Features | Expected % of Customers |
# |------|----------------|-------------|--------------|------------------------|
# | Basic | [Customer Type] | [Price] | [Features] | [Percentage] |
# | Premium | [Customer Type] | [Price] | [Features] | [Percentage] |
# | Enterprise | [Customer Type] | [Price] | [Features] | [Percentage] |

# ### 4.3 Unit Economics
# **Key Metrics:**
# - Customer Acquisition Cost (CAC): [Amount]
# - Customer Lifetime Value (LTV): [Amount]
# - LTV/CAC Ratio: [Ratio - should be >3:1]
# - Gross Margin per Customer: [Percentage]
# - Monthly Recurring Revenue (MRR) - if applicable
# - Average Revenue Per User (ARPU): [Amount]

# ---

# ## 5. GO-TO-MARKET STRATEGY & SALES

# ### 5.1 Launch Strategy
# [Collected Input - Enhanced with:]
# - Soft launch plan (beta customers, limited geography)
# - Full market launch timeline
# - Launch success metrics and KPIs
# - Marketing campaign strategy
# - PR and communications plan

# ### 5.2 Customer Acquisition Strategy
# **Customer Acquisition Channels:** [Collected Input - Enhanced with:]

# **Primary Channels:**
# 1. [Channel 1]: [Cost, conversion rate, scalability]
# 2. [Channel 2]: [Cost, conversion rate, scalability]
# 3. [Channel 3]: [Cost, conversion rate, scalability]

# **Channel Strategy by Phase:**
# - **Months 1-6:** [Focus channels and reasons]
# - **Months 7-12:** [Focus channels and reasons]
# - **Year 2+:** [Focus channels and reasons]

# **Sales Process:**
# - Lead generation strategy
# - Lead qualification criteria
# - Sales cycle length: [Duration]
# - Conversion rates at each stage
# - Sales team structure and compensation

# ### 5.3 Marketing Strategy
# **Brand Positioning:** [How you want to be perceived in the market]

# **Marketing Mix:**
# - **Product:** [Key differentiators to highlight]
# - **Price:** [Pricing strategy messaging]
# - **Place:** [Distribution and sales channels]
# - **Promotion:** [Marketing communications strategy]

# **Digital Marketing Strategy:**
# - Content marketing plan
# - Social media strategy
# - Search engine optimization (SEO)
# - Paid advertising (Google Ads, social media ads)
# - Email marketing and automation
# - Influencer and partnership marketing

# **Customer Retention Strategy:**
# - Onboarding process
# - Customer success programs
# - Loyalty programs or incentives
# - Upselling and cross-selling strategies

# ---

# ## 6. OPERATIONS PLAN

# ### 6.1 Operational Model
# **Business Operations:**
# - Key operational processes
# - Service/product delivery model
# - Quality control measures
# - Customer service approach
# - Scalability plans

# ### 6.2 Supply Chain & Vendors
# **Key Suppliers/Partners:** [If applicable]
# - Primary suppliers and backup options
# - Key partnership agreements
# - Supplier risk mitigation strategies
# - Cost management and negotiation strategies

# ### 6.3 Technology Infrastructure
# **IT Systems & Infrastructure:**
# - Core technology platforms
# - Data management and analytics
# - Cybersecurity measures
# - Disaster recovery plans
# - Scalability and performance planning

# ### 6.4 Location & Facilities
# **Physical Infrastructure:**
# - Office/facility requirements
# - Location strategy (remote, hybrid, physical)
# - Equipment and asset needs
# - Expansion plans

# ---

# ## 7. MANAGEMENT TEAM & ORGANIZATION

# ### 7.1 Founder Background
# [Collected Input - Enhanced with:]
# - Detailed professional background
# - Relevant industry experience
# - Previous entrepreneurial experience
# - Key achievements and credentials
# - Education and certifications
# - Why this founder is uniquely positioned to solve this problem

# ### 7.2 Current Team Structure
# **Planned Team & Roles:** [Collected Input - Enhanced with:]

# **Current Team:**
# - [Name]: [Title] - [Background, key skills, equity %]
# - [Name]: [Title] - [Background, key skills, equity %]

# **Organizational Chart:** [Visual representation of reporting structure]

# ### 7.3 Hiring Plan
# **Key Positions to Fill:**
# | Role | Timeline | Salary Range | Equity % | Priority | Key Qualifications |
# |------|----------|-------------|----------|----------|-------------------|
# | [Role 1] | [Month] | [Range] | [%] | [High/Medium/Low] | [Requirements] |
# | [Role 2] | [Month] | [Range] | [%] | [High/Medium/Low] | [Requirements] |

# **Team Development Strategy:**
# - Recruiting strategy and channels
# - Company culture development
# - Employee retention strategies
# - Performance management systems

# ### 7.4 Advisory Board & Mentors
# **Advisors/Partners:** [Collected Input - Enhanced with:]
# - [Name]: [Background, expertise, how they help]
# - [Name]: [Background, expertise, how they help]

# **Advisory Compensation:**
# - Typical equity grants for advisors: [0.1% - 1.0%]
# - Advisory agreement terms
# - Expected time commitment and contributions

# ### 7.5 Corporate Governance
# **Board Structure:**
# - Board composition (founders, investors, independents)
# - Board meeting frequency and format
# - Key committees (audit, compensation, etc.)
# - Decision-making processes

# ---

# ## 8. FINANCIAL PROJECTIONS & ANALYSIS

# ### 8.1 Financial Summary
# **Key Financial Highlights:**
# - Break-even timeline: [Month/Year]
# - Cash flow positive: [Month/Year]
# - Projected revenue at 3 years: [Amount]
# - Projected profit margin at maturity: [Percentage]

# ### 8.2 Revenue Projections
# **Financial Goals:** [Collected Input - Enhanced with detailed breakdown]

# | Metric | Year 1 | Year 2 | Year 3 | Year 4 | Year 5 |
# |--------|--------|--------|--------|--------|--------|
# | Revenue | [Amount] | [Amount] | [Amount] | [Amount] | [Amount] |
# | Gross Profit | [Amount] | [Amount] | [Amount] | [Amount] | [Amount] |
# | Gross Margin % | [%] | [%] | [%] | [%] | [%] |
# | EBITDA | [Amount] | [Amount] | [Amount] | [Amount] | [Amount] |
# | Net Income | [Amount] | [Amount] | [Amount] | [Amount] | [Amount] |
# | Customers | [Number] | [Number] | [Number] | [Number] | [Number] |

# ### 8.3 Expense Breakdown
# **Key Expenses:** [Collected Input - Enhanced with detailed categories]

# **Fixed Costs (Monthly):**
# - Salaries and benefits: [Amount]
# - Office rent and utilities: [Amount]
# - Software and subscriptions: [Amount]
# - Insurance: [Amount]
# - Legal and accounting: [Amount]
# - Marketing (fixed): [Amount]
# - Other fixed costs: [Amount]

# **Variable Costs (% of Revenue):**
# - Cost of goods sold: [%]
# - Sales commissions: [%]
# - Marketing (performance-based): [%]
# - Payment processing: [%]
# - Customer service: [%]

# ### 8.4 Cash Flow Analysis
# **Monthly Cash Flow Projections (Year 1):**
# [Monthly breakdown showing cash inflows, outflows, and ending balance]

# **Cash Flow Runway:**
# - Current cash on hand: [Amount]
# - Monthly burn rate: [Amount]
# - Runway without funding: [Months]
# - **Runway (How long funds will last):** [Collected Input]

# ### 8.5 Financial Assumptions
# **Key Assumptions Behind Projections:**
# - Customer acquisition rate and cost
# - Revenue per customer and growth rates
# - Pricing assumptions and changes
# - Cost inflation rates
# - Market penetration rates
# - Seasonality factors

# ### 8.6 Scenario Analysis
# **Financial Scenarios:**

# **Best Case (25% probability):**
# - Key assumptions and resulting metrics
# - Revenue and profitability outcomes

# **Base Case (50% probability):**
# - Most likely assumptions and outcomes
# - Conservative but achievable projections

# **Worst Case (25% probability):**
# - Conservative assumptions and stress testing
# - Minimum viable outcomes and survival strategies

# ---

# ## 9. FUNDING REQUEST & INVESTMENT

# ### 9.1 Funding Requirements
# **Estimated Total Budget:** [Collected Input]
# **Personal Investment:** [Collected Input - Enhanced with source of funds]
# **Expected Funding Needs:** [Collected Input]

# **Funding Round Details:**
# - **Amount Raising:** [Specific amount with range]
# - **Valuation:** [Pre-money and post-money if known]
# - **Equity Offered:** [Percentage]
# - **Type of Securities:** [Common stock, preferred stock, convertible notes]
# - **Timeline:** [Funding completion target date]

# ### 9.2 Use of Funds
# **Detailed Use of Proceeds:**
# | Category | Amount | Percentage | Timeline | Key Milestones |
# |----------|--------|------------|----------|---------------|
# | Product Development | [Amount] | [%] | [Months] | [Deliverables] |
# | Marketing & Sales | [Amount] | [%] | [Months] | [Targets] |
# | Team Expansion | [Amount] | [%] | [Months] | [Hires] |
# | Operations | [Amount] | [%] | [Months] | [Capabilities] |
# | Working Capital | [Amount] | [%] | [Months] | [Buffer] |
# | **Total** | **[Total Amount]** | **100%** | | |

# ### 9.3 Investment Returns
# **Investor Value Proposition:**
# - Expected return multiple: [X times investment]
# - Timeline to exit: [Years]
# - Exit strategy options: [IPO, acquisition, management buyout]
# - Comparable company valuations and exits

# **Key Value Drivers:**
# - Market opportunity size and growth
# - Competitive advantages and moat
# - Scalable business model
# - Experienced team
# - Clear path to profitability

# ### 9.4 Exit Strategy
# **Potential Exit Scenarios:**
# 1. **Strategic Acquisition:** [Target acquirers, timeline, valuation multiples]
# 2. **Financial Acquisition:** [Private equity, timeline, valuation multiples]
# 3. **IPO:** [Market cap requirements, timeline, precedents]
# 4. **Management Buyout:** [Conditions and structure]

# **Comparable Transactions:**
# - Recent acquisitions in the industry
# - Valuation multiples (revenue, EBITDA)
# - Strategic rationale for acquisitions

# ---

# ## 10. RISK ANALYSIS & MITIGATION

# ### 10.1 Business Risks
# **Market Risks:**
# - Market adoption slower than expected
# - Economic downturn affecting demand
# - Regulatory changes impacting the industry
# - Competitive response from established players

# **Operational Risks:**
# - Key personnel departure
# - Technology failures or security breaches
# - Supply chain disruptions
# - Scalability challenges

# **Financial Risks:**
# - Inability to raise additional funding
# - Cash flow shortfalls
# - Customer concentration risk
# - Currency or interest rate exposure

# ### 10.2 Risk Mitigation Strategies
# **For Each Major Risk:**
# - Probability assessment (High/Medium/Low)
# - Impact assessment (High/Medium/Low)
# - Specific mitigation strategies
# - Contingency plans
# - Monitoring and early warning indicators

# ### 10.3 Legal & Compliance
# **Legal/Compliance Needs:** [Collected Input - Enhanced with:]
# - Regulatory requirements and licensing
# - Intellectual property protection strategy
# - Data privacy and security compliance
# - Employment law considerations
# - Contract and liability management

# **Insurance Coverage:**
# - General liability insurance
# - Professional liability insurance
# - Cyber liability insurance
# - Key person insurance
# - Directors and officers insurance

# ---

# ## 11. MILESTONES & KEY PERFORMANCE INDICATORS

# ### 11.1 Business Milestones
# **Milestones & KPIs:** [Collected Input - Enhanced with timeline]

# **6-Month Milestones:**
# - [Milestone 1]: [Specific, measurable target]
# - [Milestone 2]: [Specific, measurable target]
# - [Milestone 3]: [Specific, measurable target]

# **12-Month Milestones:**
# - [Milestone 1]: [Specific, measurable target]
# - [Milestone 2]: [Specific, measurable target]
# - [Milestone 3]: [Specific, measurable target]

# **24-Month Milestones:**
# - [Milestone 1]: [Specific, measurable target]
# - [Milestone 2]: [Specific, measurable target]
# - [Milestone 3]: [Specific, measurable target]

# ### 11.2 Key Performance Indicators (KPIs)
# **Customer Metrics:**
# - Customer acquisition rate (customers/month)
# - Customer acquisition cost (CAC)
# - Customer lifetime value (LTV)
# - Customer retention rate
# - Net Promoter Score (NPS)
# - Monthly/Annual Recurring Revenue (MRR/ARR)

# **Financial Metrics:**
# - Revenue growth rate (month-over-month, year-over-year)
# - Gross margin percentage
# - Monthly burn rate
# - Months of runway remaining
# - Cash conversion cycle

# **Operational Metrics:**
# - Product development velocity
# - Team productivity metrics
# - Customer support response times
# - System uptime and performance

# ### 11.3 Reporting and Monitoring
# **Investor Reporting:**
# - Monthly financial reports
# - Quarterly board meetings
# - Annual strategic reviews
# - Ad-hoc updates for major developments

# **Dashboard and Analytics:**
# - Real-time KPI dashboard
# - Monthly performance reviews
# - Variance analysis against projections
# - Corrective action plans

# ---

# ## 12. ADDITIONAL CONSIDERATIONS

# ### 12.1 Environmental, Social & Governance (ESG)
# **Sustainability Initiatives:**
# - Environmental impact and sustainability measures
# - Social responsibility programs
# - Diversity and inclusion policies
# - Ethical business practices

# ### 12.2 Digital Transformation
# **Technology Adoption:**
# - Digital infrastructure and capabilities
# - Data analytics and business intelligence
# - Automation and AI integration opportunities
# - Digital customer experience enhancements

# ### 12.3 Scalability Planning
# **Growth Preparation:**
# - Systems and processes for scale
# - International expansion possibilities
# - Strategic partnership opportunities
# - Platform and ecosystem development

# ---

# ## APPENDICES

# ### Appendix A: Market Research Data
# - Industry reports and analysis
# - Customer survey results
# - Competitive intelligence
# - Market size calculations and sources

# ### Appendix B: Financial Models
# - Detailed financial projections (monthly for Year 1, quarterly for Years 2-3)
# - Sensitivity analysis
# - Scenario modeling
# - Unit economics calculations

# ### Appendix C: Technical Documentation
# - Product specifications
# - Technology architecture
# - Development roadmap
# - Intellectual property portfolio

# ### Appendix D: Legal Documents
# - Corporate structure documents
# - Key contracts and agreements
# - Regulatory compliance documentation
# - Intellectual property registrations

# ### Appendix E: Team Documentation
# - Detailed founder and team resumes
# - Organizational chart
# - Advisory board profiles
# - Recruitment and hiring plans

# ### Appendix F: Supporting Materials
# - Letters of intent from customers
# - Partnership agreements
# - Press coverage and testimonials
# - Product demos or prototypes (if available)

# ---

# ## DOCUMENT CONTROL

# **Document Version:** 1.0
# **Last Updated:** [Date]
# **Prepared By:** [Founder Name]
# **Reviewed By:** [Advisory Team]
# **Approved By:** [Board/Founders]

# **Confidentiality Notice:**
# This business plan contains confidential and proprietary information. Any reproduction or distribution of this document, in whole or in part, without written consent is strictly prohibited.

# **Contact Information:**
# [Founder Name]
# [Title]
# [Company Name]
# [Phone Number]
# [Email Address]
# [Company Address]

# ---

# *This comprehensive business plan template is designed to meet institutional investor requirements and due diligence standards. Each section should be thoroughly completed with specific, measurable data and realistic projections based on solid market research and financial analysis.*

#     Here are the contents of various documents:
# ######## Founder Checklist: ###########
#     {texts[0]}
# ######## Pitch deck: ###########
#     {texts[1]}
# ######## Additional Document 1: ###########
#     {texts[2]}
# ######## Additional Document 2: ###########
#     {texts[3]}

#     """

#     response = client.chat.completions.create(
#         model="gpt-4o-mini",
#         messages=[{"role": "user", "content": prompt}]
#     )


    # summary = response.choices[0].message.content

    summary_path = f"summaries/{request_id}.txt"
    os.makedirs("summaries", exist_ok=True)
    with open(summary_path, "w") as f:
        f.write(summary)
    # Clean up temp files
    for f in uploaded_files.values():
        os.remove(f)
    
@router.post("/analyze-documents")
async def analyze_documents(
    background_tasks: BackgroundTasks,
    request_id: str = Form(...),
    founder_name: str = Form(...),
    founder_email: str = Form(...),
    startup_name: str = Form(...),
    founderChecklist: Optional[UploadFile] = File(None),
    pitchDeck: Optional[UploadFile] = File(None),
    otherDoc1: Optional[UploadFile] = File(None),
    otherDoc2: Optional[UploadFile] = File(None),
):
    """
    Receives documents and founder/startup info, then analyzes them.
    """
    uploaded_files = dict()

    if founderChecklist:
        temp_path = UPLOAD_DIR + "/" + f"temp_{founderChecklist.filename}"
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(founderChecklist.file, buffer)
        uploaded_files["founderChecklist"] = temp_path 
    
    if pitchDeck:
        temp_path = UPLOAD_DIR + "/" +  f"temp_{pitchDeck.filename}"
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(pitchDeck.file, buffer)
        uploaded_files["pitchDeck"] = temp_path 
    
    if otherDoc1:
        temp_path = UPLOAD_DIR + "/" +  f"temp_{otherDoc1.filename}"
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(otherDoc1.file, buffer)
        uploaded_files["otherDoc1"] = temp_path 
    
    if otherDoc2:
        temp_path = UPLOAD_DIR + "/" +  f"temp_{otherDoc2.filename}"
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(otherDoc2.file, buffer)
        uploaded_files["otherDoc2"] = temp_path 

    
    # 1️⃣ Immediately send email
    message = MessageSchema(
        subject="Evalora Document Request Received",
        recipients=[founder_email],
        body=f"Your request has been submitted successfully.\nRequest ID: {request_id}",
        subtype="plain"
    )
    # fm = FastMail(conf)
    # await fm.send_message(message)

    # 2️⃣ Add background task to process documents & LLM
    # background_tasks.add_task(
    #     evaluate_startup_documents,
    #     uploaded_files,
    #     founder_name,
    #     founder_email,
    #     startup_name,
    #     request_id
    # )

     # 3️⃣ Return immediate response
    return {"message": "Request submitted successfully!", 
            "request_id": request_id,
        "founder_name": founder_name,
        "files": {
            "checklist": founderChecklist.filename if founderChecklist else None,
            "pitch": pitchDeck.filename if pitchDeck else None,
            "other1": otherDoc1.filename if otherDoc1 else None,
            "other2": otherDoc2.filename if otherDoc2 else None,
        }}

