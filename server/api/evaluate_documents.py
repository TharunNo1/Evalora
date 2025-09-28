from fastapi import Request, UploadFile, File, WebSocket, WebSocketDisconnect, APIRouter, BackgroundTasks, Form
# from fastapi.middleware.cors import CORSMiddleware
import json
from openai import OpenAI
from docx import Document
from openpyxl import Workbook
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import os
# import tempfile
from gtts import gTTS
# import base64 
from pptx import Presentation
# import fitz  # PyMuPDF
import os
import shutil
from typing import List, Optional
# from fastapi.responses import JSONResponse
# from fastapi_mail import FastMail, MessageSchema, ConnectionConfig
import base64
from utils.doc_utils import extract_text
from utils.audio_utils import get_audio_base64
from utils.session_utils import get_session, save_session, session_file
from models.gemini_client import GeminiClient
from services.gcs_service import GCSService
from services.gmail_service import GmailService
from datetime import datetime
from typing import Dict
from services.firestore_service import FirestoreService
from schemas.models import EvaluationRequest, RequestStage, Startup
import mimetypes
import re

USE_GPT_STT = False

language = 'en'
geminiClient = GeminiClient()  # Initialize GeminiClient when needed
gcsService = GCSService()
dbService = FirestoreService()

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

SESSION_DIR = "sessions"
os.makedirs(SESSION_DIR, exist_ok=True)

router = APIRouter()

gmailService = GmailService.get_instance()

# Memory of collected info per session
conversations = {}

# ---------- WebSocket voice streaming ----------
@router.post("/chat/{session_id}")
async def chat(session_id: str, request: Request):
    body = await request.json()
    user_msg = body.get("message")
    custom_prompt = body.get("prompt", '''
# 🎤 Startup Founder Voice Intake Assistant Prompt

You are a **friendly, professional voice intake assistant**. Your goal is to guide a startup founder through a conversation, collect structured details, and make them feel comfortable sharing their story.  

---

## Conversation Rules
1. Speak **conversationally** — no forms or surveys. Use natural, flowing language.  
2. Ask **one question at a time**. Wait for the founder’s spoken response before continuing.  
3. If they answer partially, ask a **follow-up or clarifying question** — don’t repeat the full original question.  
4. Encourage **storytelling**. If answers are short, politely nudge for more detail.  
5. Stay **human-like**: vary phrasing, use small acknowledgments (“Got it”, “That makes sense”, “Thanks for sharing”).  
6. After completing each section, **briefly summarize** what you captured and confirm with them.  
7. Keep responses **concise, warm, and spoken-friendly**.  

---

## Goal
Capture these structured fields:  
- **Founder & Idea Details**  
- **Market & Problem Exploration**  
- **Product & Solution Details**  
- **Operations & Team**  
- **Financials**  
- **Go-to-Market & Plan**  
- **Additional Info / Notes**  

---

## Section Prompts

### Opening / Introduction
- “Hi, thanks for taking the time to chat today. I’d love to learn more about you and your startup idea. Let’s start simple — could you tell me your name and how I can best reach you?”  
- “Great, thanks [Name]. And what’s the name of your business or startup?”  

### Founder & Idea Details
- “How would you describe your business idea in your own words?”  
- “What makes your idea stand out from what’s already out there?”  

### Market & Problem Exploration
- “Who do you imagine will benefit most from your product or service?”  
- “What’s the key problem or need you’re solving for them?”  
- “Do you have a sense of how big the market might be, or who your main competitors are?”  

### Product & Solution Details
- “Can you walk me through what your product or service actually does?”  
- “If you had to launch a simple first version, what features would it absolutely need?”  
- “What technologies or platforms do you think you’ll need to build it?”  

### Operations & Team
- “Tell me a little about your background — what led you to start this?”  
- “Do you have a team in mind already, or are you planning to bring in partners or advisors?”  

### Financials
- “What’s your rough budget estimate for getting this off the ground?”  
- “How much are you personally investing, and how much do you think you’ll need to raise?”  
- “What’s your plan for generating revenue? Subscriptions, one-time sales, something else?”  

### Go-to-Market & Plan
- “How are you thinking about launching? Soft launch, pilot, or straight to market?”  
- “What channels do you think will help you get your first customers?”  
- “Do you have milestones you want to hit in the first 6 to 12 months?”  

### Closing / Additional Notes
- “That was really helpful, thank you. Before we wrap up, is there anything else — like legal considerations, compliance issues, or special notes — that you’d like me to capture?”  
- “Perfect, thanks so much. I’ll compile everything we discussed and share it with you in a clear summary.”  

---

## Extra Realism Tips
- Vary acknowledgment phrases: “Got it”, “I see”, “That’s helpful”, “Interesting”, “Makes sense”.  
- Add **short pauses** in speech to sound natural.  
- End each section with a **mini-summary**: “So far, I have that your idea is X, targeting Y, with competitors Z. Does that sound right?”  
- If a question is skipped, move on gracefully: “No worries, we can come back to that later.”  

''')

    
    conv = get_session(session_id)

    # Store user message
    conv["messages"].append({"role": "user", "content": user_msg})

    client = OpenAI()
    # Call LLM (OpenAI example – swap model if using Mistral/LLaMA)
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": custom_prompt},
            *conv["messages"]
        ]
    )

    reply = response.choices[0].message.content
    conv["messages"].append({"role": "assistant", "content": reply})

    save_session(session_id, conv)

    speech = gTTS(text=reply, lang=language, slow=False)
    speech.save("output.mp3")

    audio_b64 = get_audio_base64()

    return {"reply": reply, "data": conv["data"], "audio_base64": audio_b64}


@router.get("/export/{session_id}/{format}")
async def export_data(session_id: str, format: str):
    
    data = get_session(session_id)
    print(data)
    if not data:
        return {"error": "No data available to export."}
    if format == "json":
        filename = "startup_details.json"
        with open(filename, "w") as f:
            json.dump(data, f, indent=4)
        return {"file": filename}

    elif format == "excel":
        filename = "startup_details.xlsx"
        wb = Workbook()
        ws = wb.active
        for section, details in data.items():
            ws.append([section])
            for k, v in details.items():
                ws.append([k, v])
            ws.append([])
        wb.save(filename)
        return {"file": filename}

    elif format == "word":
        filename = "startup_details.docx"
        doc = Document()
        for section, details in data.items():
            doc.add_heading(section, level=1)
            for k, v in details.items():
                doc.add_paragraph(f"{k}: {v}")
        doc.save(filename)
        return {"file": filename}

    elif format == "pdf":
        filename = "startup_details.pdf"
        doc = SimpleDocTemplate(filename)
        styles = getSampleStyleSheet()
        flowables = []
        for section, details in data.items():
            flowables.append(Paragraph(f"<b>{section}</b>", styles["Heading1"]))
            for k, v in details.items():
                flowables.append(Paragraph(f"{k}: {v}", styles["Normal"]))
        doc.build(flowables)
        return {"file": filename}

    return {"error": "Invalid format"}

def send_email(to, subject, body, sender="evaloraofficial@gmail.com", file_path=None):
    if file_path:
        gmailService.send_email_with_attachment(sender, to, subject, body, file_path) 
    else:
        gmailService.send_email(to, subject, body, sender)

def extract_score_from_summary(file_path: str) -> float:
    with open(file_path, "r") as f:
        content = f.read()

    match = re.search(r"EVALORA_SCORE:\s*(\d+(\.\d+)?)", content)
    if match:
        return float(match.group(1))
    return 0.0  # default if pattern not found

async def evaluate_startup_documents(background_tasks, uploaded_files, founder_name, founder_email, startup_name, request_id):
    docs = dict()

    for topic, file_paths in uploaded_files.items():
        # Ensure file_paths is always a list
        if not isinstance(file_paths, list):
            file_paths = [file_paths]

        text_combined = ""
        for file_path in file_paths:
            if file_path.endswith(".pdf"):
                text_combined += extract_text(file_path) + "\n"
            elif file_path.endswith((".ppt", ".pptx")):
                text_combined += extract_text(file_path) + "\n"
            elif file_path.endswith(".txt"):
                text_combined += extract_text(file_path) + "\n"
            elif file_path.endswith((".doc", ".docx")):
                text_combined += extract_text(file_path) + "\n"
            elif file_path.endswith((".eml")):
                text_combined += extract_text(file_path) + "\n"
            elif file_path.endswith((".mp3", ".wav", ".m4a")):
                text_combined += extract_text(file_path) + "\n"
            elif file_path.endswith(".mp4"):
                text_combined += extract_text(file_path) + "\n"
            else:
                # fallback
                text_combined += f"[Unsupported file type: {file_path}]\n"
            # os.remove(file_path)  # Clean up temp file after processing

        docs[topic] = text_combined

    
    summary = await geminiClient.analyze_documents(request_id=request_id,founder_name=founder_name,founder_email=founder_email,startup_name=startup_name,docs=docs)

    summary_path = f"summaries/{startup_name}_{request_id}.md"
    os.makedirs("summaries", exist_ok=True)
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write(summary)
    score = extract_score_from_summary(summary_path)
    is_qualified = score > 3.5

    # Send email with summary
    email_content = f"""Dear {founder_name},

Thank you for submitting your documents for your startup evaluation.

Please find attached a summary of the documents you provided.

Based on our initial review, your submission has {'met the qualification criteria' if is_qualified else 'not met the qualification criteria at this stage'}.

{"Our Evalora team will reach out to you shortly to schedule an evaluation session with an Evalora agent. During this session, we will discuss your startup in detail and provide guidance on the next steps." if is_qualified else "If you believe there are discrepancies in the evaluation, please raise a re-evaluation ticket and our team will review it."}

If you have any questions in the meantime, please feel free to contact us.

Sincerely,
Evalora Team
"""


    background_tasks.add_task(
        send_email,
        founder_email,
        f"Evalora Document Evaluation Result – {startup_name} : REQ-{request_id}",
        email_content,
        "evaloraofficial@gmail.com",
        summary_path
    )
    

    gcsService.upload_file(open(summary_path, "rb"), f"{request_id}/summary_{startup_name}_{request_id}.md", content_type="text/plain")


def upload_dynamic_file(temp_path: str, request_id: str, folder_name: str):
    # ✅ Guess MIME type from file extension
    content_type, _ = mimetypes.guess_type(temp_path)
    if content_type is None:
        content_type = "application/octet-stream"  # fallback for unknown types

    with open(temp_path, "rb") as f:
        gcsService.upload_file(
            f,
            f"{request_id}/{folder_name}/{temp_path.split('/')[-1]}",  # only filename in GCS
            content_type=content_type
        )
    
@router.post("/analyze-documents/")
async def analyze_documents(
    background_tasks: BackgroundTasks,
    request: Request,
    request_id: str = Form(...),
    description: str = Form(...),
    founder_name: str = Form(...),
    founder_email: str = Form(...),
    startup_name: str = Form(...),
    technology_subcategories: str = Form(...),   
    industry_subcategories: str = Form(...),    
    founderChecklist: Optional[UploadFile] = File(None)
):
    """
    Accepts founder/startup info and a flexible set of uploaded files:
      - founderChecklist (single)
      - pitchDeck0..n  (multiple)
      - emailMessages  (multiple)
      - callRecordings (multiple)
      - callTranscripts (multiple)
    All files are stored to disk and queued for background evaluation.
    """

    categories = []
    subcategories = {}
    tech_list = [t.strip() for t in technology_subcategories.split(",") if t.strip()]
    if len(tech_list) > 0:
        categories.append("Technology")
        subcategories["Technology"] = tech_list
    industry_list = [i.strip() for i in industry_subcategories.split(",") if i.strip()]
    if len(industry_list) > 0:
        categories.append("Industry")
        subcategories["Industry"] = industry_list
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    uploaded_files: Dict[str, List[str]] = {
        "founderChecklist": [],
        "pitchDeck": [],
        "emailMessages": [],
        "callRecordings": [],
        "callTranscripts": [],
    }

    # ---------- 1️⃣ Save checklist (if any) ----------
    if founderChecklist:
        checklist_path = os.path.join(
            UPLOAD_DIR, f"temp_{founderChecklist.filename}"
        )
        with open(checklist_path, "wb") as buffer:
            shutil.copyfileobj(founderChecklist.file, buffer)
        uploaded_files["founderChecklist"].append(checklist_path)
        upload_dynamic_file(checklist_path, request_id,"checklists")

    # ---------- 2️⃣ Save dynamic pitch decks ----------
    form = await request.form()
    for key, value in form.multi_items():
        # pitchDeck0, pitchDeck1...
        if key.startswith("pitchDeck") and isinstance(value, UploadFile):
            temp_path = os.path.join(UPLOAD_DIR, f"{value.filename}")
            with open(temp_path, "wb") as buffer:
                shutil.copyfileobj(value.file, buffer)
            gcsService.upload_file(open(temp_path, "rb"), f"{request_id}/{temp_path}", content_type="text/plain")
            uploaded_files["pitchDeck"].append(temp_path)
            upload_dynamic_file(temp_path, request_id, "pitchDecks")

    # ---------- 3️⃣ Save grouped files ----------
    async def save_multiple(prefix: str):
        files: List[UploadFile] = await request.form().then(
            lambda f: f.getlist(prefix)
        )  # FastAPI does not support getlist directly on Request
        return files

    # Since Flutter adds them via addFilesToFormData(), we must iterate manually:
    for field in ["emailMessages", "callRecordings", "callTranscripts"]:
        for key, value in form.multi_items():
            if key == field and isinstance(value, UploadFile):
                temp_path = os.path.join(UPLOAD_DIR, f"temp_{value.filename}")
                with open(temp_path, "wb") as buffer:
                    shutil.copyfileobj(value.file, buffer)
                uploaded_files[field].append(temp_path)
                upload_dynamic_file(temp_path, request_id, field)

    # ---------- 4️⃣ Background tasks ----------
    email_content = (
    f"Dear {founder_name},\n\n"
    f"We are pleased to inform you that your request has been successfully submitted.\n"
    f"Request ID: {request_id}\n\n"
    f"Our team will review the submitted documents and respond to you promptly.\n\n"
    f"If you have any questions in the meantime, please do not hesitate to contact us.\n\n"
    f"Sincerely,\n"
    f"Evalora Team"
)


    background_tasks.add_task(
        send_email,
        founder_email,
        f"Evalora Document Request – {startup_name} : REQ-{request_id}",
        email_content,
        "evaloraofficial@gmail.com",
        None
    )

    background_tasks.add_task(
        evaluate_startup_documents,
        background_tasks,
        uploaded_files,
        founder_name,
        founder_email,
        startup_name,
        request_id,
    )

    dbService.create_evaluation_request(EvaluationRequest(startupId=request_id, startupName=startup_name, description=description, founderName=founder_name, founderEmail=founder_email, docsList=uploaded_files))

    dbService.create_startup(startup=Startup(
        id=request_id,
        name=startup_name,
        description=description,
            categories=categories,
                subCategories=subcategories,
                founder=founder_name,
                founder_id="",
                score=0.0,
                currentStatus=RequestStage.submission,
                approved=False
                ))

    return {
        "message": "Request submitted successfully!",
        "request_id": request_id,
        "founder_name": founder_name,
        "uploaded": {k: [os.path.basename(p) for p in v]
                     for k, v in uploaded_files.items()},
    }
